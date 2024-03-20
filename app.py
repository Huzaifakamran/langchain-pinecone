import streamlit as st
import os
from docx import Document
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone
# from langchain_openai import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.memory import ConversationBufferWindowMemory, ConversationBufferMemory
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.embeddings import OpenAIEmbeddings
from langchain.llms import OpenAI
from dotenv import load_dotenv
load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
os.environ["OPENAI_API_KEY"] = api_key
index_name = os.getenv("PINECONE_INDEX_NAME")
embeddings = OpenAIEmbeddings()
text_field = "text"

vectorstore = PineconeVectorStore(
    index_name, embeddings, text_field
)

st.session_state.filesData = ''
class Documents:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        return f"Document(page_content='{self.page_content}', metadata={self.metadata})"


template = """You are an AI ASSISTANT that gives information related to the provided documents 

Context: {context}

History: {history}

Question: {question}
# """

prompt = PromptTemplate(
    input_variables=["history","context","question"],
    template=template,
)

def query_db(query):
    query = query
    global qa_chain
    llm_response = qa_chain(query)
    # response = process_llm_response(llm_response)
    # print(response)
    return llm_response

def set_model(vectordb):
    pcIndex2 = pc.Index(index_name)
    vectorstore2 = PineconeVectorStore(pcIndex2, embeddings,"text")
    retriever = vectorstore2.as_retriever()
    
    # create the chain to answer questions 
    global qa_chain
    # completion llm
    qa_chain = RetrievalQA.from_chain_type(
                llm=ChatOpenAI(model_name = 'gpt-4', temperature=0),
                chain_type='stuff',
                retriever=retriever,
                verbose=True,
                chain_type_kwargs={
                    "verbose": True,
                    "prompt": prompt,
                    "memory": ConversationBufferWindowMemory(
                        memory_key="history",
                        k=30,
                        input_key="question"),
                }
            )
    return qa_chain

def setUpPINECONEDB():
    # configure client
    
    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    return pc
    # pinecone.init(
    #     api_key=os.getenv("PINECONE_API_KEY"),
    #     environment=os.getenv("PINECONE_ENVIRONMENT")
    # )
pc = setUpPINECONEDB()

    
def insertDocs(doc,embeddings,index_name):
    index = PineconeVectorStore.from_documents(
        doc, index_name=index_name, embedding=embeddings
    )
    return index
    
# Function to convert data to desired format
def convert_to_documents(data):
    documents = []
    for entry in data:
        page_content = entry['page_content']
        metadata = {}  # Replace with actual metadata extraction

        # Extract relevant metadata from the provided entry
        # For demonstration purposes, only 'source' is extracted
        metadata['source'] = entry['metadata']['source']

        # Append Document object to the list
        documents.append(Documents(page_content, metadata))

    return documents

def shape_document(documents,filename):
    
    # one_pera = ''
    # documents_content = [doc.page_content for doc in documents]
    # for doc in documents_content:
        # one_pera += doc
        # one_pera += ' '
    # print(st.session_state.text_file_content)
    # if st.session_state.text_file_content:
    #     one_pera += st.session_state.text_file_content
    one_pera = [documents]
    # documents_content
    # metadata = [doc.metadata for doc in documents]
    # metadatas = [metadata[i]['source'] for i in range(len(metadata))]
    data = []
    for i in range(len(one_pera)):
        data.append({'page_content': one_pera[i], 'metadata': {'source': filename}})
    document_objects = convert_to_documents(data)
    # print(document_objects)
    return document_objects

def read_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_pdf(file):
    with open(file, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

def setUpUi():
    pcIndex = pc.Index(index_name)
    normal_index = PineconeVectorStore(pcIndex, embeddings)
    st.session_state.index = normal_index
    qa_chain = set_model(vectorstore)
    st.session_state.qa_chain = qa_chain
    st.sidebar.header("Upload Documents")
    st.sidebar.write("Select the documents you want to upload:")
    uploaded_files = st.sidebar.file_uploader("Choose files", accept_multiple_files=True)

    if uploaded_files:
        st.sidebar.write("You have uploaded the following files:")
        for file in uploaded_files:
            st.sidebar.write(file.name)

    if st.sidebar.button("Upload Documets"):
        
        
        
        document_objects = []
        if uploaded_files:
            for file in uploaded_files:
                with open(file.name, "wb") as f:
                    f.write(file.getbuffer())
                    
                file_extension = os.path.splitext(file.name)[1]
                if file_extension == ".docx":
                    text = read_docx(file.name)
                elif file_extension == ".pdf":
                    text = read_pdf(file.name)
                else:
                    text = "Document format not supported for direct display."
                document_objects.append(shape_document(text,file.name)[0])
                # print(document_objects)
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
            texts = text_splitter.split_documents(document_objects)
            
            index = insertDocs(texts,embeddings,index_name)
            st.session_state.index = index


# Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
    else:
        # Display chat messages from history on app rerun
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                # React to user input
    if prompt := st.chat_input("What is up?"):
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(prompt)
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.chat_message("assistant"):
            print(st.session_state.qa_chain)
            response_bot = st.session_state.qa_chain({"query": prompt})
            # print metadata source of response
            print(response_bot)
            st.markdown(response_bot['result'])
            st.session_state.chat_history.append((prompt, response_bot['result']))
        # Add assistant response to chat history
        st.session_state.messages.append({"role": "assistant", "content": response_bot['result']})
    # st.write(st.session_state.messages)    

setUpUi()
